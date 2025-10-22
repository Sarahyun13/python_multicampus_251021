from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_kfood_document():
    # 새로운 Document 객체 생성
    doc = Document()
    
    # 제목 추가
    title = doc.add_heading('한국의 대표적인 음식(K-Food)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 소개 문단 추가
    doc.add_paragraph('K-Food는 건강하고 다양한 맛으로 전 세계적으로 인기를 얻고 있습니다.')
    
    # 대표 음식 리스트
    popular_foods = [
        {
            'name': '김치',
            'description': '한국의 대표적인 발효 음식으로, 항산화 물질이 풍부하고 건강에 좋습니다.',
            'popularity': '전 세계적으로 인정받는 건강식품'
        },
        {
            'name': '비빔밥',
            'description': '다양한 채소와 고기를 고추장과 함께 비벼먹는 영양만점 한 그릇 식사입니다.',
            'popularity': '건강식으로 해외에서 큰 인기'
        },
        {
            'name': '불고기',
            'description': '달콤한 양념의 구운 쇠고기 요리로, 외국인들이 가장 좋아하는 한식 중 하나입니다.',
            'popularity': '한국 요리의 대표 주자'
        }
    ]
    
    # 음식 정보 추가
    for food in popular_foods:
        heading = doc.add_heading(food['name'], level=1)
        doc.add_paragraph(food['description'])
        p = doc.add_paragraph('인기도: ')
        p.add_run(food['popularity']).bold = True
        doc.add_paragraph('\n')
    
    # 마무리 문단
    doc.add_paragraph('K-Food는 건강하고 맛있는 식문화를 전파하며 계속해서 세계적으로 성장하고 있습니다.')
    
    # 문서 저장
    doc.save('k-food.docx')

if __name__ == '__main__':
    create_kfood_document()
    print("k-food.docx 파일이 생성되었습니다.")